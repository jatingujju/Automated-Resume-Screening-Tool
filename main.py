print("🔥 Script started")

import sqlite3
import os
import json
import uuid
import pandas as pd

from PyPDF2 import PdfReader
import docx

from src.extract import parse_resume
from src.features import jd_resume_features


# -----------------------------
# EXTRACT TEXT
# -----------------------------
def extract_text(file_path):
    try:
        if file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            return "".join([p.extract_text() or "" for p in reader.pages])

        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            return " ".join([p.text for p in doc.paragraphs])

        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

        return ""

    except Exception as e:
        print(f"❌ Error reading {file_path}: {e}")
        return ""


# -----------------------------
# LOAD JOB DESCRIPTION
# -----------------------------
def load_jd(path):
    if not os.path.exists(path):
        print("❌ Job description file not found!")
        return ""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


# -----------------------------
# MAIN
# -----------------------------
# EXTRACT TEXT
# -----------------------------
def extract_text(file_path):
    try:
        if file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            return "".join([p.extract_text() or "" for p in reader.pages])

        elif file_path.endswith(".docx"):
            doc = docx.Document(file_path)
            return " ".join([p.text for p in doc.paragraphs])

        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

        return ""

    except Exception as e:
        print(f"❌ Error reading {file_path}: {e}")
        return ""
# -----------------------------
if __name__ == "__main__":

    print("🚀 Running Resume Screening System...\n")

    # Load JD
    print("➡️ Loading Job Description...")
    jd_text = load_jd("data/job_description.txt")

    if not jd_text:
        print("❌ JD is empty. Exiting.")
        exit()

    print("✅ JD Loaded\n")

    # Load resumes
    print("➡️ Checking resumes folder...")
    if not os.path.exists("resumes"):
        print("❌ 'resumes' folder not found!")
        exit()

    files = [f for f in os.listdir("resumes") if f.endswith((".pdf", ".docx", ".txt"))]

    if len(files) == 0:
        print("❌ No valid resume files found!")
        exit()

    print("📁 Found Resumes:", files)

    results = []
    job_id = str(uuid.uuid4())

    # -----------------------------
    # PROCESS EACH RESUME
    # -----------------------------
    for file in files:
        resume_path = os.path.join("resumes", file)

        print(f"\n📄 Processing: {file}")

        # Extract
        print("➡️ Extracting text...")
        raw = extract_text(resume_path)

        if not raw.strip():
            print("⚠️ Empty or unreadable file, skipping...")
            continue

        print("✅ Text extracted | Length:", len(raw))

        # Parse
        parsed = parse_resume(raw)

        print("📊 Parsed Data:", parsed)

        # Save to DB
        candidate_id = str(uuid.uuid4())

        conn = sqlite3.connect("data/resume.db")
        cursor = conn.cursor()

        cursor.execute("""
            INSERT OR REPLACE INTO resumes (candidate_id, source, raw_text, parsed_json)
            VALUES (?, ?, ?, ?)
        """, (
            candidate_id,
            "file",
            raw,
            json.dumps(parsed)
        ))

        cursor.execute("""
            INSERT OR IGNORE INTO jobs (id, title, jd_text, must_have)
            VALUES (?, ?, ?, ?)
        """, (
            job_id,
            "Demo Job",
            jd_text,
            "communication, video editing, teamwork, leadership"
        ))

        conn.commit()
        conn.close()

        # Features
        features = jd_resume_features("data/resume.db", job_id, candidate_id)

        print("📊 Features:", features)

        # -----------------------------
        # SCORING (IMPROVED)
        # -----------------------------
        if features:
            skill_ratio = features["rule_musthave_hits"] / max(1, features["rule_musthave_total"])

            # 👇 fresher-friendly experience score
            exp_score = min(features["years_exp"] / 5, 1) if features["years_exp"] else 0.3

            final_score = (
                0.3 * features["sim_embedding"] +
                0.5 * skill_ratio +
                0.2 * exp_score
            )
        else:
            final_score = 0

        print(f"🎯 Final Score: {round(final_score * 100, 2)}%")

        # Store
        results.append({
            "Candidate": file,
            "Skills": ", ".join(parsed.get("skills", [])),
            "Score (%)": round(final_score * 100, 2)
        })

    # -----------------------------
    # FINAL RANKING
    # -----------------------------
    if not results:
        print("\n❌ No valid results.")
        exit()

    df = pd.DataFrame(results).sort_values(by="Score (%)", ascending=False)

    print("\n🏆 FINAL RANKING:\n")
    print(df)

    os.makedirs("outputs", exist_ok=True)
    df.to_csv("outputs/ranking.csv", index=False)

    print("\n✅ Results saved to outputs/ranking.csv")